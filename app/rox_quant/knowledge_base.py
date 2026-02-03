from dataclasses import dataclass, field
from typing import List, Optional, Any
import os
import json
import re
import traceback

HAS_NLP = False
cosine_similarity = None

@dataclass
class KnowledgeDocument:
    path: str
    title: str
    content: str
    vector: Optional[List[float]] = field(default=None)

class KnowledgeBase:
    def __init__(self):
        self.documents: List[KnowledgeDocument] = []
        self.model = None
        self._nlp_checked = False

    def _ensure_nlp(self) -> bool:
        global HAS_NLP
        global cosine_similarity
        if self._nlp_checked:
            return bool(HAS_NLP and self.model and cosine_similarity)
        self._nlp_checked = True
        try:
            from sentence_transformers import SentenceTransformer
            from sklearn.metrics.pairwise import cosine_similarity as _cos
            cosine_similarity = _cos
            HAS_NLP = True
            self.model = SentenceTransformer("all-MiniLM-L6-v2")
            return True
        except ImportError:
            HAS_NLP = False
            self.model = None
            return False
        except Exception as e:
            print(f"Failed to load SentenceTransformer: {e}")
            HAS_NLP = False
            self.model = None
            return False

    def _normalize_text(self, s: str) -> str:
        lines = [re.sub(r"\s+", " ", ln).strip() for ln in s.splitlines()]
        lines = [ln for ln in lines if ln]
        return "\n".join(lines)

    def _parse_pdf(self, fp: str) -> str:
        content = ""
        try:
            from pdfminer.high_level import extract_text as _pdf_extract  # type: ignore
            content = _pdf_extract(fp) or ""
        except Exception:
            content = ""
        if not content:
            try:
                import fitz  # type: ignore
                d = fitz.open(fp)
                content = "\n".join([p.get_text() or "" for p in d])
            except Exception:
                content = ""
        if not content:
            try:
                from PyPDF2 import PdfReader  # type: ignore
                r = PdfReader(fp)
                content = "\n".join([(page.extract_text() or "") for page in r.pages])
            except Exception:
                content = ""
        return self._normalize_text(content or "")
    
    def _compute_embedding(self, text: str) -> Optional[List[float]]:
        if not self._ensure_nlp():
            return None
        try:
            # 截断过长的文本以适应模型限制 (通常 256 或 512 tokens)
            # 简单按字符截断，例如前 1000 字符
            embedding = self.model.encode(text[:1000])
            return embedding.tolist()
        except Exception as e:
            print(f"Embedding error: {e}")
            return None

    def build_embedded_from_dir(self, src_dir: str) -> int:
        arr = []
        if not os.path.isdir(src_dir):
            return 0
        
        print(f"Building KB from {src_dir}...")
        for root, _, files in os.walk(src_dir):
            for f in files:
                if f.startswith("~$") or f.startswith("."):
                    continue
                fp = os.path.join(root, f)
                ext = os.path.splitext(fp)[1].lower()
                title = os.path.splitext(os.path.basename(fp))[0]
                content = ""
                try:
                    if ext in [".txt", ".md"]:
                        with open(fp, "r", encoding="utf-8", errors="ignore") as fh:
                            content = self._normalize_text(fh.read())
                    elif ext == ".docx":
                        try:
                            from docx import Document  # type: ignore
                            d = Document(fp)
                            content = self._normalize_text("\n".join([p.text for p in d.paragraphs]))
                        except Exception:
                            content = ""
                    elif ext == ".pdf":
                        content = self._parse_pdf(fp)
                    else:
                        content = ""
                except Exception:
                    content = ""
                
                if content:
                    # 生成向量
                    vec = self._compute_embedding(title + "\n" + content)
                    arr.append({"path": fp, "title": title, "content": content, "vector": vec})
                    print(f"Processed: {title}")

        base = os.path.dirname(__file__)
        assets_dir = os.path.join(base, "assets")
        if not os.path.exists(assets_dir):
            os.makedirs(assets_dir)
            
        out = os.path.join(assets_dir, "embedded_kb.json")
        try:
            with open(out, "w", encoding="utf-8") as f:
                json.dump(arr, f, ensure_ascii=False)
            print(f"Saved {len(arr)} docs to {out}")
        except Exception as e:
            print(f"Failed to save KB: {e}")
            return 0
        return len(arr)

    def load_embedded(self) -> int:
        self.documents.clear()
        # Use resource_utils if available, otherwise fallback
        try:
            from ..resource_utils import get_resource_path
            path = get_resource_path(os.path.join("app", "rox_quant", "assets", "embedded_kb.json"))
        except ImportError:
             base = os.path.dirname(__file__)
             path = os.path.join(base, "assets", "embedded_kb.json")

        if not os.path.isfile(path):
            # Try direct relative path if resource path fails or returns non-existent
             base = os.path.dirname(__file__)
             local_path = os.path.join(base, "assets", "embedded_kb.json")
             if os.path.isfile(local_path):
                 path = local_path
             else:
                 return 0

        try:
            with open(path, "r", encoding="utf-8") as f:
                arr = json.load(f)
            if isinstance(arr, list):
                for item in arr:
                    t = str(item.get("title") or "").strip()
                    c = str(item.get("content") or "").strip()
                    p = str(item.get("path") or "")
                    v = item.get("vector") # List[float] or None
                    if t and c:
                        self.documents.append(KnowledgeDocument(path=p, title=t, content=self._normalize_text(c), vector=v))
        except Exception:
            self.documents = []
        return len(self.documents)

    def load_dir(self, path: str) -> int:
        self.documents.clear()
        if not os.path.isdir(path):
            return 0
        for root, _, files in os.walk(path):
            for f in files:
                if f.startswith("~$") or f.startswith("."):
                    continue
                fp = os.path.join(root, f)
                ext = os.path.splitext(fp)[1].lower()
                title = os.path.splitext(os.path.basename(fp))[0]
                content = ""
                try:
                    if ext in [".txt", ".md"]:
                        with open(fp, "r", encoding="utf-8", errors="ignore") as fh:
                            content = self._normalize_text(fh.read())
                    elif ext == ".docx":
                        try:
                            from docx import Document  # type: ignore
                            d = Document(fp)
                            content = self._normalize_text("\n".join([p.text for p in d.paragraphs]))
                        except Exception:
                            content = ""
                    elif ext == ".pdf":
                        content = self._parse_pdf(fp)
                    else:
                        content = ""
                except Exception:
                    content = ""
                if content:
                    vec = self._compute_embedding(title + "\n" + content)
                    self.documents.append(KnowledgeDocument(path=fp, title=title, content=content, vector=vec))
        return len(self.documents)

    def size(self) -> int:
        return len(self.documents)

    def search(self, query: str, limit: int = 5) -> List[KnowledgeDocument]:
        # 混合搜索：如果有向量则结合语义，否则仅关键词
        q = query.strip().lower()
        if not q:
            return []
        
        # 1. 语义搜索 (Semantic Search)
        semantic_scores = {} # id -> score
        if self._ensure_nlp():
            try:
                q_vec = self.model.encode(q)
                # 收集所有有向量的文档
                valid_docs = [(i, d.vector) for i, d in enumerate(self.documents) if d.vector]
                if valid_docs:
                    ids = [x[0] for x in valid_docs]
                    vecs = [x[1] for x in valid_docs]
                    
                    # 计算相似度
                    sims = cosine_similarity([q_vec], vecs)[0]
                    
                    for idx, score in zip(ids, sims):
                        semantic_scores[idx] = float(score)
            except Exception as e:
                print(f"Semantic search error: {e}")

        # 2. 关键词搜索 (Keyword Search)
        keyword_scores = {}
        for i, doc in enumerate(self.documents):
            s = 0
            if q in doc.title.lower():
                s += 3.0
            if q in doc.content.lower():
                s += 1.0
            # 简单的词频统计
            s += doc.content.lower().count(q) * 0.1
            if s > 0:
                keyword_scores[i] = s

        # 3. 融合分数 (Fusion)
        # 归一化关键词分数 (简单的 max-min 归一化)
        if keyword_scores:
            max_kw = max(keyword_scores.values())
            if max_kw > 0:
                for k in keyword_scores:
                    keyword_scores[k] /= max_kw # map to 0-1

        final_scores = []
        for i, doc in enumerate(self.documents):
            sem_s = semantic_scores.get(i, 0.0)
            kw_s = keyword_scores.get(i, 0.0)
            
            # 加权融合: 语义 0.7 + 关键词 0.3 (可调整)
            # 如果没有语义库，全靠关键词
            if HAS_NLP and self.model:
                total = sem_s * 0.7 + kw_s * 0.3
            else:
                total = kw_s
            
            if total > 0.01: # 阈值
                final_scores.append((total, doc))

        final_scores.sort(key=lambda x: x[0], reverse=True)
        return [d for _, d in final_scores[:limit]]

    def get_context_for_algo(self, topic: str) -> str:
        """
        为算法模块提供特定主题的上下文知识
        """
        docs = self.search(topic, limit=3)
        if not docs:
            return ""
        
        context = []
        for d in docs:
            # 截取摘要
            snippet = d.content[:500].replace("\n", " ")
            context.append(f"《{d.title}》: {snippet}...")
        return "\n".join(context)

    def count_sector_keywords(self) -> dict:
        keywords = {
            "新能源": 0,
            "半导体": 0,
            "人工智能": 0,
            "数据资产": 0,
            "光伏": 0,
            "电动车": 0,
            "医药": 0,
            "消费": 0,
            "白酒": 0,
            "低空经济": 0,
            "华为": 0,
            "算力": 0
        }
        for doc in self.documents:
            text = doc.content.lower()
            for k in list(keywords.keys()):
                keywords[k] += text.count(k.lower())
        return keywords

    def count_macro_keywords(self) -> dict:
        keywords = {
            "加息": 0,
            "降息": 0,
            "通胀": 0,
            "GDP": 0,
            "宽货币": 0,
            "紧信用": 0,
            "政策支持": 0,
            "风险": 0,
            "复苏": 0,
            "流动性": 0
        }
        for doc in self.documents:
            text = doc.content.lower()
            for k in list(keywords.keys()):
                keywords[k] += text.count(k.lower())
        return keywords

