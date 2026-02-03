"""
Word文档读取工具 - 提取文本、表格和图片
"""

import os
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
import json

def extract_docx_content(docx_path, output_dir='./docx_extract'):
    """
    完整提取Word文档的所有内容
    
    包括：
    - 文本内容
    - 表格
    - 图片（走势图等）
    """
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    
    doc = Document(docx_path)
    
    # 1. 提取基本信息
    print("\n" + "="*80)
    print("📖 Word文档内容提取")
    print("="*80)
    print(f"✓ 总段落数: {len(doc.paragraphs)}")
    print(f"✓ 总表格数: {len(doc.tables)}")
    print(f"✓ 输出目录: {output_dir}")
    
    # 2. 提取所有文本
    print("\n" + "-"*80)
    print("【文本内容】")
    print("-"*80)
    
    text_content = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            text_content.append(para.text)
            print(f"{i+1:3}. {para.text[:100]}")
    
    # 3. 提取所有表格
    print("\n" + "-"*80)
    print("【表格内容】")
    print("-"*80)
    
    tables_data = []
    for table_idx, table in enumerate(doc.tables):
        print(f"\n表格 {table_idx + 1}:")
        table_rows = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_rows.append(row_data)
            print(f"  {row_data}")
        tables_data.append(table_rows)
    
    # 4. 提取所有图片
    print("\n" + "-"*80)
    print("【图片提取】")
    print("-"*80)
    
    image_count = 0
    image_info = []
    
    # 方法1: 从段落中提取图片
    for para_idx, para in enumerate(doc.paragraphs):
        for run in para.runs:
            for rel in run._element.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'):
                try:
                    image_rId = rel.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    image_part = doc.part.related_part(image_rId)
                    image_bytes = image_part.blob
                    
                    # 获取图片格式
                    image_ext = image_part.partname.split('.')[-1]
                    image_filename = f"image_paragraph_{para_idx}_{image_count}.{image_ext}"
                    image_path = os.path.join(output_dir, image_filename)
                    
                    with open(image_path, 'wb') as f:
                        f.write(image_bytes)
                    
                    print(f"✓ 从段落{para_idx}提取图片: {image_filename}")
                    image_info.append({
                        'source': f'paragraph_{para_idx}',
                        'filename': image_filename,
                        'size': len(image_bytes),
                        'format': image_ext
                    })
                    image_count += 1
                except Exception as e:
                    pass
    
    # 方法2: 从表格中提取图片
    for table_idx, table in enumerate(doc.tables):
        for cell_idx, cell in enumerate(table._cells):
            for rel in cell._element.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'):
                try:
                    image_rId = rel.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    image_part = doc.part.related_part(image_rId)
                    image_bytes = image_part.blob
                    
                    image_ext = image_part.partname.split('.')[-1]
                    image_filename = f"image_table_{table_idx}_cell_{cell_idx}_{image_count}.{image_ext}"
                    image_path = os.path.join(output_dir, image_filename)
                    
                    with open(image_path, 'wb') as f:
                        f.write(image_bytes)
                    
                    print(f"✓ 从表格{table_idx}提取图片: {image_filename}")
                    image_info.append({
                        'source': f'table_{table_idx}',
                        'filename': image_filename,
                        'size': len(image_bytes),
                        'format': image_ext
                    })
                    image_count += 1
                except Exception as e:
                    pass
    
    # 5. 保存为JSON格式
    print("\n" + "-"*80)
    print("【保存为JSON】")
    print("-"*80)
    
    json_data = {
        'document_info': {
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'total_images': image_count
        },
        'text_content': text_content,
        'tables': tables_data,
        'images': image_info
    }
    
    json_path = os.path.join(output_dir, 'document_content.json')
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)
    
    print(f"✓ JSON已保存: {json_path}")
    
    # 6. 保存纯文本
    text_path = os.path.join(output_dir, 'document_text.txt')
    with open(text_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(text_content))
    
    print(f"✓ 文本已保存: {text_path}")
    
    # 7. 总结
    print("\n" + "="*80)
    print("✅ 提取完成!")
    print("="*80)
    print(f"📊 统计:")
    print(f"  • 文本段落: {len(text_content)}")
    print(f"  • 表格: {len(tables_data)}")
    print(f"  • 图片: {image_count}")
    print(f"\n📁 输出文件:")
    print(f"  • {json_path}")
    print(f"  • {text_path}")
    if image_count > 0:
        print(f"  • {output_dir}/image_*.* (图片文件)")
    
    return json_data


if __name__ == '__main__':
    # 修改这个路径为你的Word文档路径
    docx_path = '/Users/mac/Downloads/rox3.0/app/data/量化交易从入门到精通 - 未知.docx'
    
    # 检查文件是否存在
    if not os.path.exists(docx_path):
        print(f"❌ 文件不存在: {docx_path}")
        print("\n尝试查找相似文件...")
        import subprocess
        result = subprocess.run(
            ['find', '/Users/mac/Downloads', '-name', '*量化*', '-type', 'f'],
            capture_output=True, text=True
        )
        if result.stdout:
            print("找到的相似文件:")
            for line in result.stdout.strip().split('\n')[:10]:
                print(f"  {line}")
        exit(1)
    
    # 提取内容
    content = extract_docx_content(docx_path, output_dir='./docx_extract')
    
    # 显示部分内容预览
    print("\n" + "="*80)
    print("📄 内容预览 (前500个字)")
    print("="*80)
    full_text = '\n'.join(content['text_content'])
    print(full_text[:500])
