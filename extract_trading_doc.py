#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Extract content from Word document: 量化交易从入门到精通 - 未知.docx
"""
import sys
from docx import Document

def extract_content():
    try:
        doc = Document('app/data/量化交易从入门到精通 - 未知.docx')
        
        print("=" * 80)
        print(f"文档统计: {len(doc.paragraphs)} 段落, {len(doc.tables)} 表格")
        print("=" * 80)
        
        # 提取所有段落
        print("\n【所有段落内容】\n")
        para_count = 0
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                para_count += 1
                print(f"{para_count}. {text}")
                if para_count > 50:  # 只显示前50段
                    print(f"\n... 还有 {len(doc.paragraphs) - i - 1} 段 ...")
                    break
        
        # 提取表格
        print("\n\n【表格内容】\n")
        for table_idx, table in enumerate(doc.tables):
            print(f"\n表格 {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                print(" | ".join(cells))
                if row_idx > 20:
                    print("...")
                    break
        
        print("\n\n【文档提取完成】")
        
    except Exception as e:
        print(f"错误: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    extract_content()
