#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import os

# 路径定义
doc_path = "/Users/mac/Downloads/rox3.0/app/data/量化交易从入门到精通 - 未知.docx"
output_path = "/Users/mac/Downloads/rox3.0/extracted_trading_knowledge.txt"

# 检查文件是否存在
if not os.path.exists(doc_path):
    print(f"Error: File not found - {doc_path}")
    exit(1)

# 加载Word文档
doc = Document(doc_path)

# 初始化数据收集
paragraphs_list = []
tables_info = []
headings = []
key_principles = []

# 提取段落
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        paragraphs_list.append(text)
        if para.style.name.startswith('Heading'):
            headings.append(text)

# 提取表格
for table_idx, table in enumerate(doc.tables, 1):
    table_data = {'index': table_idx, 'rows': []}
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        table_data['rows'].append(row_data)
    tables_info.append(table_data)

# 识别关键交易原则
keywords = ['原则', '规则', '指标', '策略', '风险', '交易', '买入', '卖出', '支撑', '阻力', '趋势']
for para in paragraphs_list:
    if any(keyword in para for keyword in keywords):
        if len(para) < 200:
            key_principles.append(para)

# 生成输出文件
with open(output_path, 'w', encoding='utf-8') as f:
    f.write("=" * 80 + "\n")
    f.write("量化交易知识提取结果\n")
    f.write("=" * 80 + "\n\n")
    
    # 文档统计
    f.write("【文档统计】\n")
    f.write("-" * 40 + "\n")
    f.write("段落总数: " + str(len(paragraphs_list)) + "\n")
    f.write("表格总数: " + str(len(tables_info)) + "\n")
    f.write("标题总数: " + str(len(headings)) + "\n")
    f.write("识别的交易原则: " + str(len(key_principles)) + "\n")
    f.write("\n")
    
    # 完整内容
    f.write("【完整内容】\n")
    f.write("-" * 40 + "\n")
    for idx, para in enumerate(paragraphs_list, 1):
        f.write(str(idx) + ". " + para + "\n\n")
    
    f.write("\n")
    
    # 表格信息
    f.write("【表格信息】\n")
    f.write("-" * 40 + "\n")
    for table_data in tables_info:
        f.write("\n表格 " + str(table_data['index']) + ":\n")
        for row_idx, row_data in enumerate(table_data['rows'], 1):
            f.write("  行 " + str(row_idx) + ": " + " | ".join(row_data) + "\n")
    
    f.write("\n")
    
    # 关键信息总结
    f.write("【关键信息总结】\n")
    f.write("-" * 40 + "\n")
    if key_principles:
        for idx, principle in enumerate(key_principles, 1):
            f.write(str(idx) + ". " + principle + "\n\n")
    else:
        f.write("未识别到具体的交易原则关键词。\n")
    
    # 标题列表
    if headings:
        f.write("\n【文档标题结构】\n")
        f.write("-" * 40 + "\n")
        for heading in headings:
            f.write("- " + heading + "\n")

# 计算内容字符总数
total_chars = sum(len(para) for para in paragraphs_list)
for table_data in tables_info:
    for row in table_data['rows']:
        total_chars += sum(len(cell) for cell in row)

# 返回结果
print("Successfully extracted content")
print("Paragraphs: " + str(len(paragraphs_list)))
print("Tables: " + str(len(tables_info)))
print("Total chars: " + str(total_chars))
print("Key principles: " + str(len(key_principles)))
print("Output file: " + output_path)
print("File size: " + str(os.path.getsize(output_path)) + " bytes")
