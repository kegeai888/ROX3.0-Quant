#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从Word文档中提取交易知识
"""
import sys
import os
sys.path.insert(0, '/Users/mac/Downloads/rox3.0')

try:
    from docx import Document
    
    doc_path = '/Users/mac/Downloads/rox3.0/app/data/量化交易从入门到精通 - 未知.docx'
    
    if not os.path.exists(doc_path):
        print(f"文件不存在: {doc_path}")
        sys.exit(1)
    
    doc = Document(doc_path)
    
    # 提取所有文本
    all_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text)
    
    # 保存到文件
    output_file = '/Users/mac/Downloads/rox3.0/doc_content.txt'
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write(f"从Word文档提取的交易知识\n")
        f.write("=" * 80 + "\n\n")
        f.write(f"总段落数: {len(doc.paragraphs)}\n")
        f.write(f"有内容段落数: {len(all_text)}\n")
        f.write(f"表格数: {len(doc.tables)}\n\n")
        
        # 写入所有文本
        for i, text in enumerate(all_text, 1):
            f.write(f"{i}. {text}\n")
        
        # 写入表格
        if doc.tables:
            f.write("\n\n" + "=" * 80 + "\n")
            f.write("表格内容\n")
            f.write("=" * 80 + "\n\n")
            for table_idx, table in enumerate(doc.tables, 1):
                f.write(f"\n【表格 {table_idx}】\n")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    f.write(" | ".join(cells) + "\n")
    
    print(f"成功提取文档内容到: {output_file}")
    print(f"总字数: {sum(len(t) for t in all_text)}")
    print(f"\n前10段内容:")
    for i, text in enumerate(all_text[:10], 1):
        print(f"{i}. {text[:100]}...")

except ImportError as e:
    print(f"缺少依赖: {e}")
    print("运行: pip install python-docx")
except Exception as e:
    print(f"错误: {e}")
    import traceback
    traceback.print_exc()
