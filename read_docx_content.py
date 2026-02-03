#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
读取Word文档：量化交易从入门到精通
提取所有文本、表格和图片
"""

from docx import Document
import os
import json

def extract_docx():
    docx_path = '/Users/mac/Downloads/rox3.0/app/data/量化交易从入门到精通 - 未知.docx'
    
    print("🔍 检查文件...")
    if not os.path.exists(docx_path):
        print(f"❌ 文件不存在: {docx_path}")
        return
    
    print(f"✓ 文件存在: {os.path.getsize(docx_path)} 字节")
    
    try:
        doc = Document(docx_path)
        print("✓ 文档已加载")
    except Exception as e:
        print(f"❌ 加载失败: {e}")
        return
    
    # 提取基本信息
    print(f"\n📊 文档统计:")
    print(f"  • 段落数: {len(doc.paragraphs)}")
    print(f"  • 表格数: {len(doc.tables)}")
    
    # 提取文本内容
    print(f"\n📄 文本内容 (前30段):")
    print("="*80)
    
    text_content = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            text_content.append(para.text)
            if i < 30:
                # 限制每行长度显示
                text = para.text[:100]
                print(f"{i+1:3}. {text}...")
    
    print(f"\n✓ 提取了 {len(text_content)} 段非空文本")
    
    # 提取表格
    print(f"\n📋 表格内容:")
    print("="*80)
    
    tables_info = []
    for table_idx, table in enumerate(doc.tables):
        print(f"\n【表格 {table_idx + 1}】({len(table.rows)} 行 x {len(table.columns)} 列)")
        rows_data = []
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text for cell in row.cells]
            rows_data.append(row_data)
            if row_idx < 5:  # 只显示前5行
                print(f"  行{row_idx+1}: {row_data}")
        tables_info.append(rows_data)
    
    # 提取图片
    print(f"\n🖼️ 图片提取:")
    print("="*80)
    
    image_count = 0
    image_list = []
    
    # 从段落提取
    for para_idx, para in enumerate(doc.paragraphs):
        for run in para.runs:
            # 查找所有关系（可能包含图片）
            try:
                for rel_key in run._element.getparent().getparent().part.rels:
                    pass
            except:
                pass
            
            # 查找嵌入的图片
            try:
                for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    # 找到图片引用
                    embed = drawing.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}blip')
                    if embed is not None:
                        try:
                            rId = embed.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if rId:
                                image_part = doc.part.related_part(rId)
                                image_bytes = image_part.blob
                                image_ext = image_part.partname.split('.')[-1]
                                
                                image_filename = f"chart_{image_count}.{image_ext}"
                                image_list.append({
                                    'source': f'paragraph_{para_idx}',
                                    'filename': image_filename,
                                    'size': len(image_bytes),
                                    'format': image_ext
                                })
                                
                                print(f"✓ 图片 {image_count}: {image_filename} ({len(image_bytes)} 字节)")
                                image_count += 1
                        except:
                            pass
            except:
                pass
    
    # 从表格提取
    for table_idx, table in enumerate(doc.tables):
        for cell_idx, cell in enumerate(table._cells):
            try:
                for drawing in cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    embed = drawing.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}blip')
                    if embed is not None:
                        try:
                            rId = embed.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if rId:
                                image_part = doc.part.related_part(rId)
                                image_bytes = image_part.blob
                                image_ext = image_part.partname.split('.')[-1]
                                
                                image_filename = f"table_chart_{table_idx}_{image_count}.{image_ext}"
                                image_list.append({
                                    'source': f'table_{table_idx}',
                                    'filename': image_filename,
                                    'size': len(image_bytes),
                                    'format': image_ext
                                })
                                
                                print(f"✓ 表格图片 {image_count}: {image_filename} ({len(image_bytes)} 字节)")
                                image_count += 1
                        except:
                            pass
            except:
                pass
    
    # 保存统计信息
    print(f"\n📊 最终统计:")
    print("="*80)
    print(f"  ✓ 文本段落: {len(text_content)}")
    print(f"  ✓ 表格: {len(tables_info)}")
    print(f"  ✓ 图片: {image_count}")
    
    # 保存JSON
    output = {
        'summary': {
            'total_paragraphs': len(doc.paragraphs),
            'total_text_paragraphs': len(text_content),
            'total_tables': len(tables_info),
            'total_images': image_count
        },
        'text_content': text_content[:50],  # 前50段
        'tables_count': [len(t) for t in tables_info],
        'images': image_list
    }
    
    json_path = '/Users/mac/Downloads/rox3.0/docx_content.json'
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    
    print(f"\n✅ 已保存JSON到: {json_path}")
    
    # 显示部分内容
    print(f"\n📖 文本内容概览 (前300字):")
    print("="*80)
    full_text = '\n'.join(text_content)
    print(full_text[:300])
    
    return {
        'text_count': len(text_content),
        'table_count': len(tables_info),
        'image_count': image_count
    }

if __name__ == '__main__':
    result = extract_docx()
    if result:
        print(f"\n✅ 提取完成！")
        print(f"   文本: {result['text_count']}")
        print(f"   表格: {result['table_count']}")
        print(f"   图片: {result['image_count']}")
