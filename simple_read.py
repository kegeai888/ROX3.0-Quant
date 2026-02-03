#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

docx_file = '/Users/mac/Downloads/rox3.0/app/data/量化交易从入门到精通 - 未知.docx'

print("开始读取Word文档...")
doc = Document(docx_file)

print(f"\n段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

print("\n【文本内容 - 前30段】")
print("="*80)

count = 0
for para in doc.paragraphs:
    if para.text.strip() and count < 30:
        print(f"{count+1}. {para.text}")
        count += 1

print(f"\n【表格内容】")
print("="*80)

for table_idx, table in enumerate(doc.tables):
    print(f"\n表格 {table_idx+1}:")
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text for cell in row.cells]
        print(f"  行{row_idx+1}: {cells}")

print("\n✅ 读取完成！")
